import streamlit as st
from supabase import create_client
import datetime
from groq import Groq
import json
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="CLC Team Meetings", page_icon="👥", layout="wide")

# ── Supabase ──────────────────────────────────────────────────────────────────
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
ADMIN_PASSWORD = st.secrets.get("ADMIN_PASSWORD", "CLC2026admin")
GROQ_KEY = st.secrets.get("GROQ_API_KEY", "")

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

# ── Helpers ───────────────────────────────────────────────────────────────────
PROGRAMS = {
    "JP":    {"label": "Junior Primary",  "color": "#2d7d4f", "light": "#d1fae5", "emoji": "🟢"},
    "PY":    {"label": "Primary Years",   "color": "#1a4d8c", "light": "#dbeafe", "emoji": "🔵"},
    "SY":    {"label": "Senior Years",    "color": "#7c3aed", "light": "#ede9fe", "emoji": "🟣"},
    "STAFF": {"label": "Staff Meetings",  "color": "#1e6f75", "light": "#ccfbf1", "emoji": "👥"},
    "ADMIN": {"label": "Admin Meetings",  "color": "#92400e", "light": "#fef3c7", "emoji": "🏢"},
}

DAYS_OF_WEEK = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
RECURRENCE_TYPES = ["Weekly", "Fortnightly", "Monthly"]

def get_program_from_params():
    params = st.query_params
    prog = params.get("program", "JP")
    return prog if prog in PROGRAMS else "JP"

def get_week_from_params():
    """Return (monday, sunday) date tuple if ?week= param present, else None."""
    week_str = st.query_params.get("week", None)
    if not week_str:
        return None
    try:
        monday = datetime.date.fromisoformat(week_str)
        monday = monday - datetime.timedelta(days=monday.weekday())
        return monday, monday + datetime.timedelta(days=6)
    except ValueError:
        return None

def admin_login():
    if "admin" not in st.session_state:
        st.session_state.admin = False
    with st.sidebar:
        if not st.session_state.admin:
            with st.expander("🔐 Admin Login"):
                pw = st.text_input("Password", type="password", key="pw_input")
                if st.button("Login"):
                    if pw == ADMIN_PASSWORD:
                        st.session_state.admin = True
                        st.rerun()
                    else:
                        st.error("Incorrect password")
        else:
            st.success("✅ Admin logged in")
            if st.button("Logout"):
                st.session_state.admin = False
                st.rerun()

def improve_with_ai(raw_text: str, program: str, meeting_date: str) -> str:
    """Use Groq to improve meeting minutes from raw transcript."""
    if not GROQ_KEY:
        return raw_text
    client = Groq(api_key=GROQ_KEY)
    prompt = f"""You are a professional minute-taker for an alternative education school (Cowandilla Learning Centre, Learning & Behaviour Unit).

Program team: {PROGRAMS[program]['label']} ({program})
Meeting date: {meeting_date}

Transform the following raw transcript or rough notes into polished, professional meeting minutes.

Format requirements:
- Clear heading with meeting details
- Numbered agenda items with concise summaries
- Action items clearly identified with: Action | Assigned to | Timeline
- Professional but accessible language
- Trauma-informed, strengths-based framing where relevant
- Preserve all factual content — do not add or fabricate information
- End with a summary table of all action items

Raw input:
{raw_text}"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=4000
    )
    return response.choices[0].message.content

# ── DB helpers ────────────────────────────────────────────────────────────────
def fetch(table, filters=None, order=None):
    q = supabase.table(table).select("*")
    if filters:
        for k, v in filters.items():
            q = q.eq(k, v)
    if order:
        q = q.order(order, desc=True)
    return q.execute().data or []

def insert(table, data):
    supabase.table(table).insert(data).execute()

def update_row(table, row_id, data):
    supabase.table(table).update(data).eq("id", row_id).execute()

def delete_row(table, row_id):
    supabase.table(table).delete().eq("id", row_id).execute()

# ─────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT GENERATION
# ─────────────────────────────────────────────────────────────────────────────



# ─────────────────────────────────────────────────────────────────────────────
# SECTION RENDERERS — SCHEDULES & AGENDA
# ─────────────────────────────────────────────────────────────────────────────

def render_schedules(program):
    """Meeting schedules — recurring + additional one-offs."""
    st.subheader("📅 Meeting Schedules")

    schedules = fetch("tm_schedules", {"program": program}, "created_at")

    if schedules:
        for s in schedules:
            rec_label = s.get("recurrence_type", "Weekly")
            day_label = s.get("day_of_week", "")
            time_label = s.get("meeting_time", "")
            flags = []
            if s.get("show_in_calendar"): flags.append("📅 Calendar")
            if s.get("show_in_bulletin"): flags.append("📋 Bulletin")
            flag_str = "  ·  ".join(flags) if flags else "No flags"

            with st.expander(
                f"📌 {s.get('schedule_name','Unnamed')} — {rec_label} {day_label} {time_label}",
                expanded=False
            ):
                c1, c2, c3 = st.columns(3)
                c1.markdown(f"**Recurrence:** {rec_label}")
                c2.markdown(f"**Day/Time:** {day_label} {time_label}")
                c3.markdown(f"**Flags:** {flag_str}")
                if s.get("notes"):
                    st.caption(s["notes"])
                if st.session_state.get("admin"):
                    col_e, col_d = st.columns(2)
                    with col_d:
                        if st.button("🗑️ Delete", key=f"del_sched_{s['id']}"):
                            delete_row("tm_schedules", s["id"])
                            st.rerun()
    else:
        st.info("No schedules set up yet.")

    if st.session_state.get("admin"):
        with st.expander("➕ Add Schedule"):
            with st.form(f"sched_form_{program}"):
                s_name = st.text_input("Schedule name *", placeholder="e.g. Weekly JP Team Meeting")
                sc1, sc2, sc3 = st.columns(3)
                with sc1:
                    s_rec  = st.selectbox("Recurrence", RECURRENCE_TYPES)
                with sc2:
                    s_day  = st.selectbox("Day of week", DAYS_OF_WEEK)
                with sc3:
                    s_time = st.text_input("Time", placeholder="e.g. 9:00am")
                s_cal  = st.checkbox("Show in Calendar")
                s_bull = st.checkbox("Show in Bulletin")
                s_notes = st.text_area("Notes (optional)", height=60)
                if st.form_submit_button("Add Schedule"):
                    if s_name.strip():
                        insert("tm_schedules", {
                            "program":          program,
                            "schedule_name":    s_name.strip(),
                            "recurrence_type":  s_rec,
                            "day_of_week":      s_day,
                            "meeting_time":     s_time.strip(),
                            "show_in_calendar": s_cal,
                            "show_in_bulletin": s_bull,
                            "notes":            s_notes.strip(),
                        })
                        st.success("Schedule added!")
                        st.rerun()
                    else:
                        st.warning("Please enter a schedule name.")


def render_agenda(program):
    """Agenda items — any staff can submit, admin manages status."""
    st.subheader("📋 Agenda")

    items = fetch("tm_agenda_items", {"program": program}, "created_at")

    open_items  = [i for i in items if i.get("status","open") != "closed"]
    closed_items = [i for i in items if i.get("status","open") == "closed"]

    if open_items:
        for item in open_items:
            c1, c2 = st.columns([5, 1])
            with c1:
                st.markdown(f"📌 **{item['title']}**")
                if item.get("description"):
                    st.caption(item["description"])
                st.caption(f"Submitted by {item.get('submitted_by','—')}  ·  {str(item.get('created_at',''))[:10]}")
            with c2:
                if st.session_state.get("admin"):
                    if st.button("✅ Close", key=f"close_agenda_{item['id']}"):
                        update_row("tm_agenda_items", item["id"], {"status": "closed"})
                        st.rerun()
                    if st.button("🗑️", key=f"del_agenda_{item['id']}"):
                        delete_row("tm_agenda_items", item["id"])
                        st.rerun()
            st.divider()
    else:
        st.info("No open agenda items.")

    if closed_items:
        with st.expander(f"✅ Closed items ({len(closed_items)})"):
            for item in closed_items:
                st.markdown(f"~~{item['title']}~~ — {item.get('submitted_by','—')}")

    with st.expander("➕ Submit an Agenda Item"):
        with st.form(f"agenda_form_{program}"):
            a_title = st.text_input("Item title *")
            a_desc  = st.text_area("Description (optional)", height=80)
            a_by    = st.text_input("Your name")
            if st.form_submit_button("Submit"):
                if a_title.strip():
                    insert("tm_agenda_items", {
                        "program":      program,
                        "title":        a_title.strip(),
                        "description":  a_desc.strip(),
                        "submitted_by": a_by.strip(),
                        "status":       "open",
                    })
                    st.success("Agenda item submitted!")
                    st.rerun()
                else:
                    st.warning("Please enter a title.")


# ─────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def _hex(colour):
    """Ensure hex string is clean (no #, uppercase)."""
    return colour.lstrip("#").upper()

def _shade_cell(cell, hex_colour):
    try:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  _hex(hex_colour))
        tcPr.append(shd)
    except Exception:
        pass

def _rgb(hex_colour):
    h = _hex(hex_colour)
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _para_run(para, text, bold=False, italic=False, size_pt=10, colour="0f172a"):
    run = para.add_run(str(text))
    run.bold         = bold
    run.italic       = italic
    run.font.size    = Pt(size_pt)
    run.font.color.rgb = _rgb(colour)
    return run

def _add_banner(doc, text, bg="1e293b", fg="FFFFFF", size=12):
    """Full-width banner paragraph acting as a section separator."""
    table = doc.add_table(rows=1, cols=1)
    cell  = table.cell(0, 0)
    _shade_cell(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    _para_run(p, f"  {text}", bold=True, size_pt=size, colour=fg)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def _add_kv_table(doc, rows, col1_bg="1e3a5f", col1_fg="FFFFFF", col2_bg="f8fafc", col2_fg="0f172a"):
    """Key/value two-column table."""
    table = doc.add_table(rows=len(rows), cols=2)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, (k, v) in enumerate(rows):
        c1, c2 = table.cell(i, 0), table.cell(i, 1)
        if i == 0:
            _shade_cell(c1, col1_bg); _shade_cell(c2, col1_bg)
            _para_run(c1.paragraphs[0], k, bold=True, size_pt=9, colour=col1_fg)
            _para_run(c2.paragraphs[0], v, bold=True, size_pt=9, colour=col1_fg)
        else:
            _shade_cell(c1, "e2e8f0"); _shade_cell(c2, col2_bg)
            _para_run(c1.paragraphs[0], k, bold=True, size_pt=9, colour="334155")
            _para_run(c2.paragraphs[0], v, bold=False, size_pt=9, colour=col2_fg)
    try:
        for row in table.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(12)
    except Exception:
        pass
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _add_attendance_table(doc, present, apologies):
    """Two-column present / apologies table."""
    rows = max(len(present), len(apologies), 1)
    table = doc.add_table(rows=rows + 1, cols=2)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for ci, (hdr, bg) in enumerate([("✅  Present", "2d7d4f"), ("⚠️  Apologies", "92400e")]):
        c = table.cell(0, ci)
        _shade_cell(c, bg)
        _para_run(c.paragraphs[0], hdr, bold=True, size_pt=9, colour="FFFFFF")
    for ri in range(rows):
        pn = present[ri]   if ri < len(present)   else ""
        an = apologies[ri] if ri < len(apologies) else ""
        for ci, name in enumerate([pn, an]):
            cell = table.cell(ri + 1, ci)
            _shade_cell(cell, "f8fafc" if ri % 2 == 0 else "FFFFFF")
            _para_run(cell.paragraphs[0], name, size_pt=9)
    try:
        for row in table.rows:
            row.cells[0].width = Cm(8.5)
            row.cells[1].width = Cm(8.5)
    except Exception:
        pass
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _add_digital_items_table(doc, items):
    """Digital meeting items table: Item | Raised By | Actioned By | Status."""
    if not items:
        p = doc.add_paragraph("No digital meeting items recorded.")
        p.paragraph_format.space_after = Pt(6)
        return
    headers = ["Item / Notice", "Raised By", "Actioned By", "Status"]
    table = doc.add_table(rows=len(items) + 1, cols=4)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for ci, hdr in enumerate(headers):
        c = table.cell(0, ci)
        _shade_cell(c, "1e3a5f")
        _para_run(c.paragraphs[0], hdr, bold=True, size_pt=9, colour="FFFFFF")
    for ri, item in enumerate(items):
        status = str(item.get("status", "Noted"))
        row_bg = "f0f9ff" if ri % 2 == 0 else "FFFFFF"
        status_bg = "d1fae5" if "action" in status.lower() else "fef3c7" if "pending" in status.lower() else "f1f5f9"
        vals = [
            item.get("item", ""),
            item.get("raised_by", ""),
            item.get("actioned_by", ""),
            status,
        ]
        for ci, val in enumerate(vals):
            cell = table.cell(ri + 1, ci)
            _shade_cell(cell, status_bg if ci == 3 else row_bg)
            _para_run(cell.paragraphs[0], str(val), size_pt=9)
    widths = [8.5, 3, 3, 2.5]
    try:
        for ci, w in enumerate(widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    except Exception:
        pass
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _add_actions_table(doc, action_text):
    """Action items: Action | Assigned To | Due Date (pipe-delimited lines)."""
    lines = [l.strip() for l in action_text.strip().splitlines() if l.strip()]
    if not lines:
        return
    _add_banner(doc, "✅  ACTION ITEMS SUMMARY", bg="4c1d95", fg="FFFFFF", size=11)
    table = doc.add_table(rows=len(lines) + 1, cols=3)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for ci, hdr in enumerate(["Action", "Assigned To", "Due Date"]):
        c = table.cell(0, ci)
        _shade_cell(c, "5b21b6")
        _para_run(c.paragraphs[0], hdr, bold=True, size_pt=9, colour="FFFFFF")
    for ri, line in enumerate(lines):
        parts = [p.strip() for p in line.split("|")]
        for ci in range(3):
            cell = table.cell(ri + 1, ci)
            _shade_cell(cell, "f5f3ff" if ri % 2 == 0 else "FFFFFF")
            _para_run(cell.paragraphs[0], parts[ci] if ci < len(parts) else "", size_pt=9)
    try:
        for row in table.rows:
            row.cells[0].width = Cm(9.5)
            row.cells[1].width = Cm(4)
            row.cells[2].width = Cm(3.5)
    except Exception:
        pass
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _add_content_block(doc, text, placeholder="Not recorded."):
    """Add body text, preserving line breaks."""
    content = (text or "").strip()
    if not content:
        p = doc.add_paragraph(placeholder)
        p.paragraph_format.space_after = Pt(4)
        return
    for line in content.splitlines():
        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(2)
            p.runs[0].font.size = Pt(10) if p.runs else None
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _doc_cover(doc, program, title, meeting_date):
    """Dark cover header block."""
    prog = PROGRAMS[program]
    prog_colours = {
        "JP": "2d7d4f", "PY": "1a4d8c", "SY": "7c3aed",
        "STAFF": "1e6f75", "ADMIN": "92400e"
    }
    accent = prog_colours.get(program, "1e293b")

    table = doc.add_table(rows=1, cols=1)
    cell  = table.cell(0, 0)
    _shade_cell(cell, "0f172a")

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after  = Pt(2)
    _para_run(p1, "  Cowandilla Learning Centre — Learning & Behaviour Unit",
              bold=True, size_pt=9, colour="94a3b8")

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(2)
    _para_run(p2, f"  {title or prog['label'] + ' Meeting Minutes'}",
              bold=True, size_pt=17, colour="FFFFFF")

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after  = Pt(8)
    _para_run(p3, f"  {prog['label']}   ·   {meeting_date}",
              bold=False, size_pt=10, colour=accent.upper() if len(accent) == 6 else "7dd3fc")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def generate_combined_docx(m: dict) -> bytes:
    """
    Generate Word doc for a STAFF combined meeting (digital + face-to-face).
    m is the raw Supabase row dict.
    """
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin   = Cm(2.2)
            section.right_margin  = Cm(2.2)

        prog = "STAFF"
        _doc_cover(doc, prog, m.get("title",""), m.get("meeting_date",""))

        # ── Meeting details ───────────────────────────────────────────────────
        _add_banner(doc, "📋  MEETING DETAILS", bg="1e3a5f")
        _add_kv_table(doc, [
            ("Field",    "Details"),
            ("Program",  PROGRAMS[prog]["label"]),
            ("Date",     m.get("meeting_date", "—")),
            ("Chair",    m.get("chair", "—") or "—"),
            ("Location", m.get("location", "—") or "—"),
        ])

        # ── Attendance ────────────────────────────────────────────────────────
        _add_banner(doc, "🙋  ATTENDANCE", bg="1e3a5f")
        present   = [n.strip() for n in (m.get("attendees")  or "").split(",") if n.strip()]
        apologies = [n.strip() for n in (m.get("apologies")  or "").split(",") if n.strip()]
        _add_attendance_table(doc, present, apologies)

        # ── Part 1 Digital ────────────────────────────────────────────────────
        _add_banner(doc, "💻  PART 1 — DIGITAL STAFF MEETING", bg="1a4d8c")

        digital_summary = (m.get("digital_summary") or "").strip()
        if digital_summary:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            _para_run(p, "Digital Meeting Summary", bold=True, size_pt=10, colour="1a4d8c")
            _add_content_block(doc, digital_summary)

        p = doc.add_paragraph()
        _para_run(p, "Items Raised — Viewed & Actioned Status", bold=True, size_pt=10, colour="1a4d8c")
        p.paragraph_format.space_after = Pt(4)

        try:
            digital_items = json.loads(m.get("digital_items") or "[]")
        except Exception:
            digital_items = []
        _add_digital_items_table(doc, digital_items)

        # ── Part 2 Face-to-face ───────────────────────────────────────────────
        _add_banner(doc, "👥  PART 2 — FACE-TO-FACE MEETING", bg="2d7d4f")
        ff = (m.get("face_to_face_content") or m.get("content") or "").strip()
        _add_content_block(doc, ff, placeholder="No face-to-face minutes recorded.")

        # ── Actions ───────────────────────────────────────────────────────────
        action_text = (m.get("action_summary") or "").strip()
        if action_text:
            _add_actions_table(doc, action_text)

        # ── Footer ────────────────────────────────────────────────────────────
        doc.add_paragraph()
        footer_table = doc.add_table(rows=1, cols=1)
        fc = footer_table.cell(0, 0)
        _shade_cell(fc, "f1f5f9")
        fp = fc.paragraphs[0]
        fp.paragraph_format.space_before = Pt(4)
        fp.paragraph_format.space_after  = Pt(4)
        _para_run(fp,
            f"  Generated {datetime.datetime.now().strftime('%d %b %Y, %I:%M %p')}  "
            f"·  CLC Learning & Behaviour Unit  ·  CONFIDENTIAL",
            italic=True, size_pt=8, colour="64748b")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        # Fallback: plain document with error note
        doc2 = Document()
        doc2.add_paragraph(f"Error generating formatted document: {e}")
        doc2.add_paragraph(f"Title: {m.get('title','')}")
        doc2.add_paragraph(f"Date: {m.get('meeting_date','')}")
        doc2.add_paragraph(f"Chair: {m.get('chair','')}")
        doc2.add_paragraph(f"Attendees: {m.get('attendees','')}")
        doc2.add_paragraph(f"Apologies: {m.get('apologies','')}")
        doc2.add_paragraph("--- DIGITAL MEETING ---")
        doc2.add_paragraph(m.get("digital_summary",""))
        doc2.add_paragraph("--- FACE TO FACE ---")
        doc2.add_paragraph(m.get("face_to_face_content","") or m.get("content",""))
        doc2.add_paragraph("--- ACTIONS ---")
        doc2.add_paragraph(m.get("action_summary",""))
        buf2 = io.BytesIO()
        doc2.save(buf2)
        buf2.seek(0)
        return buf2.getvalue()


def generate_team_docx(m: dict, program: str) -> bytes:
    """
    Generate Word doc for JP/PY/SY team meeting minutes.
    m is the raw Supabase row dict.
    """
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin   = Cm(2.2)
            section.right_margin  = Cm(2.2)

        prog_colours = {
            "JP": "2d7d4f", "PY": "1a4d8c", "SY": "7c3aed",
            "STAFF": "1e6f75", "ADMIN": "92400e"
        }
        col = prog_colours.get(program, "1e293b")

        _doc_cover(doc, program, m.get("title",""), m.get("meeting_date",""))

        _add_banner(doc, "📋  MEETING DETAILS", bg=col)
        _add_kv_table(doc, [
            ("Field",    "Details"),
            ("Program",  PROGRAMS[program]["label"]),
            ("Date",     m.get("meeting_date", "—")),
            ("Chair",    m.get("chair", "—") or "—"),
            ("Location", m.get("location", "—") or "—"),
        ])

        _add_banner(doc, "🙋  ATTENDANCE", bg=col)
        present   = [n.strip() for n in (m.get("attendees")  or "").split(",") if n.strip()]
        apologies = [n.strip() for n in (m.get("apologies")  or "").split(",") if n.strip()]
        _add_attendance_table(doc, present, apologies)

        _add_banner(doc, "📝  MEETING MINUTES", bg=col)
        content = (m.get("content") or m.get("face_to_face_content") or "").strip()
        _add_content_block(doc, content, placeholder="No minutes recorded.")

        action_text = (m.get("action_summary") or "").strip()
        if action_text:
            _add_actions_table(doc, action_text)

        doc.add_paragraph()
        footer_table = doc.add_table(rows=1, cols=1)
        fc = footer_table.cell(0, 0)
        _shade_cell(fc, "f1f5f9")
        fp = fc.paragraphs[0]
        fp.paragraph_format.space_before = Pt(4)
        fp.paragraph_format.space_after  = Pt(4)
        _para_run(fp,
            f"  Generated {datetime.datetime.now().strftime('%d %b %Y, %I:%M %p')}  "
            f"·  CLC Learning & Behaviour Unit  ·  CONFIDENTIAL",
            italic=True, size_pt=8, colour="64748b")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        doc2 = Document()
        doc2.add_paragraph(f"Error generating formatted document: {e}")
        doc2.add_paragraph(f"Title: {m.get('title','')}")
        doc2.add_paragraph(f"Date: {m.get('meeting_date','')}")
        doc2.add_paragraph(f"Chair: {m.get('chair','')}")
        doc2.add_paragraph(f"Attendees: {m.get('attendees','')}")
        doc2.add_paragraph(m.get("content",""))
        doc2.add_paragraph(m.get("action_summary",""))
        buf2 = io.BytesIO()
        doc2.save(buf2)
        buf2.seek(0)
        return buf2.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# SECTION RENDERERS
# ─────────────────────────────────────────────────────────────────────────────

def render_minutes(program, week_range=None):
    """Minutes — unified view with download for all programs."""

    # Programme accent colours for banners
    prog_colours = {
        "JP": "#2d7d4f", "PY": "#1a4d8c", "SY": "#7c3aed",
        "STAFF": "#1e6f75", "ADMIN": "#92400e"
    }
    accent = prog_colours.get(program, "#1e293b")
    is_staff = (program == "STAFF")

    # Linked-week notice
    if week_range:
        monday, sunday = week_range
        week_str = f"{monday.strftime('%-d %b')} – {sunday.strftime('%-d %b %Y')}"
        st.markdown(f"""
        <div style="background:#e8edf3;border:1px solid #b8cfe8;border-radius:8px;
                    padding:10px 16px;margin-bottom:14px;font-size:13px;color:#1a2e44;">
          🔗 <strong>Linked from Digital Staff Meeting</strong> — week of {week_str}
        </div>""", unsafe_allow_html=True)

    minutes_list = fetch("tm_minutes", {"program": program}, "meeting_date")

    # ── VIEW EXISTING ─────────────────────────────────────────────────────────
    if minutes_list:
        for m in minutes_list:
            auto_expand = False
            if week_range:
                try:
                    md = datetime.date.fromisoformat(m["meeting_date"])
                    auto_expand = week_range[0] <= md <= week_range[1]
                except Exception:
                    pass

            label = f"📄  {m.get('meeting_date','')}  —  {m.get('title','Untitled')}"
            with st.expander(label, expanded=auto_expand):

                # ── META STRIP ────────────────────────────────────────────────
                st.markdown(f"""
                <div style="background:{accent};border-radius:8px;padding:10px 16px;
                            margin-bottom:12px;display:flex;gap:2rem;flex-wrap:wrap;">
                  <span style="color:white;font-size:13px;">
                    <strong>Chair:</strong>&nbsp;{m.get('chair','—') or '—'}
                  </span>
                  <span style="color:white;font-size:13px;">
                    <strong>Location:</strong>&nbsp;{m.get('location','—') or '—'}
                  </span>
                  <span style="color:white;font-size:13px;">
                    <strong>Date:</strong>&nbsp;{m.get('meeting_date','—')}
                  </span>
                </div>""", unsafe_allow_html=True)

                # ── ATTENDANCE ────────────────────────────────────────────────
                present_raw   = m.get("attendees", "") or ""
                apologies_raw = m.get("apologies", "") or ""
                present_list   = [n.strip() for n in present_raw.split(",")   if n.strip()]
                apologies_list = [n.strip() for n in apologies_raw.split(",") if n.strip()]

                if present_list or apologies_list:
                    st.markdown("""
                    <div style="background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;
                                padding:10px 16px;margin-bottom:12px;">
                      <strong style="font-size:13px;">🙋 Attendance</strong>
                    </div>""", unsafe_allow_html=True)
                    ac1, ac2 = st.columns(2)
                    with ac1:
                        st.markdown("**✅ Present**")
                        st.markdown("\n".join(f"- {n}" for n in present_list) if present_list else "_None recorded_")
                    with ac2:
                        st.markdown("**⚠️ Apologies**")
                        st.markdown("\n".join(f"- {n}" for n in apologies_list) if apologies_list else "_None_")

                # ── DIGITAL SECTION (STAFF) ───────────────────────────────────
                if is_staff:
                    st.markdown("""
                    <div style="background:#dbeafe;border-left:5px solid #1a4d8c;
                                border-radius:0 8px 8px 0;padding:10px 16px;margin:14px 0 8px;">
                      <strong style="color:#1a4d8c;font-size:14px;">
                        💻 Part 1 — Digital Staff Meeting
                      </strong>
                    </div>""", unsafe_allow_html=True)

                    digital_summary = (m.get("digital_summary") or "").strip()
                    if digital_summary:
                        st.markdown(digital_summary)
                    else:
                        st.caption("_No digital meeting summary recorded._")

                    # Items table
                    st.markdown("**Items raised — viewed & actioned:**")
                    try:
                        items = json.loads(m.get("digital_items") or "[]")
                    except Exception:
                        items = []

                    if items:
                        import pandas as pd
                        df = pd.DataFrame(items)
                        df.columns = [c.replace("_", " ").title() for c in df.columns]
                        st.dataframe(df, use_container_width=True, hide_index=True)
                    else:
                        st.caption("_No items recorded._")

                    # ── FACE TO FACE SECTION ──────────────────────────────────
                    st.markdown("""
                    <div style="background:#d1fae5;border-left:5px solid #2d7d4f;
                                border-radius:0 8px 8px 0;padding:10px 16px;margin:14px 0 8px;">
                      <strong style="color:#2d7d4f;font-size:14px;">
                        👥 Part 2 — Face-to-Face Meeting
                      </strong>
                    </div>""", unsafe_allow_html=True)

                    ff = (m.get("face_to_face_content") or m.get("content") or "").strip()
                    if ff:
                        st.markdown(ff)
                    else:
                        st.caption("_No face-to-face minutes recorded._")

                else:
                    # JP/PY/SY
                    content = (m.get("content") or m.get("face_to_face_content") or "").strip()
                    if content:
                        st.markdown(content)
                    else:
                        st.caption("_No minutes content recorded._")

                # ── ACTION SUMMARY ────────────────────────────────────────────
                action_text = (m.get("action_summary") or "").strip()
                if action_text:
                    st.markdown("""
                    <div style="background:#f5f3ff;border-left:5px solid #5b21b6;
                                border-radius:0 8px 8px 0;padding:8px 14px;margin:12px 0 4px;">
                      <strong style="color:#5b21b6;">✅ Action Items</strong>
                    </div>""", unsafe_allow_html=True)
                    lines = [l.strip() for l in action_text.splitlines() if l.strip()]
                    for line in lines:
                        parts = [p.strip() for p in line.split("|")]
                        if len(parts) >= 2:
                            st.markdown(f"🔲 **{parts[0]}** — {parts[1]}" + (f" _(due {parts[2]})_" if len(parts) > 2 else ""))
                        else:
                            st.markdown(f"🔲 {line}")

                # ── DOWNLOAD ──────────────────────────────────────────────────
                st.markdown("")
                if is_staff:
                    docx_bytes = generate_combined_docx(m)
                else:
                    docx_bytes = generate_team_docx(m, program)

                safe_title = (m.get("title") or "minutes").replace(" ", "_")
                st.download_button(
                    label     = "📥 Download Word Document",
                    data      = docx_bytes,
                    file_name = f"{safe_title}_{m.get('meeting_date','')}.docx",
                    mime      = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key       = f"dl_{m['id']}",
                    type      = "primary",
                )

                # ── ADMIN EDIT / DELETE ───────────────────────────────────────
                if st.session_state.get("admin"):
                    st.divider()
                    ec1, ec2 = st.columns(2)
                    with ec1:
                        if st.button("✏️ Edit", key=f"edit_min_{m['id']}"):
                            st.session_state[f"editing_min_{m['id']}"] = True
                    with ec2:
                        if st.button("🗑️ Delete", key=f"del_min_{m['id']}"):
                            delete_row("tm_minutes", m["id"])
                            st.rerun()
                    if st.session_state.get(f"editing_min_{m['id']}"):
                        with st.form(f"edit_min_form_{m['id']}"):
                            ec1, ec2 = st.columns([3, 1])
                            with ec1: etitle = st.text_input("Title", value=m.get("title",""))
                            with ec2: edate  = st.date_input("Date", value=datetime.date.fromisoformat(m["meeting_date"]) if m.get("meeting_date") else datetime.date.today())
                            ec3, ec4, ec5 = st.columns(3)
                            with ec3: echair = st.text_input("Chair", value=m.get("chair",""))
                            with ec4: eloc   = st.text_input("Location", value=m.get("location",""))
                            with ec5: pass
                            eatt  = st.text_input("Present (comma-separated)", value=m.get("attendees",""))
                            eapol = st.text_input("Apologies (comma-separated)", value=m.get("apologies",""))
                            if is_staff:
                                edigital = st.text_area("Digital meeting summary", value=m.get("digital_summary",""), height=120)
                                eitems   = st.text_area("Digital items (Item|Raised By|Actioned By|Status)", value="\n".join(
                                    f"{i.get('item','')}|{i.get('raised_by','')}|{i.get('actioned_by','')}|{i.get('status','')}"
                                    for i in json.loads(m.get("digital_items") or "[]")
                                ), height=100)
                                eff      = st.text_area("Face-to-face minutes", value=m.get("face_to_face_content",""), height=200)
                                econtent = ""
                            else:
                                edigital = ""; eitems = ""; eff = ""
                                econtent = st.text_area("Minutes", value=m.get("content",""), height=280)
                            eactions = st.text_area("Action items (Action|Assigned To|Due)", value=m.get("action_summary",""))
                            if st.form_submit_button("💾 Save Changes"):
                                parsed_items = []
                                if is_staff:
                                    for line in eitems.splitlines():
                                        parts = [p.strip() for p in line.split("|")]
                                        if parts and parts[0]:
                                            parsed_items.append({
                                                "item": parts[0], "raised_by": parts[1] if len(parts)>1 else "",
                                                "actioned_by": parts[2] if len(parts)>2 else "",
                                                "status": parts[3] if len(parts)>3 else "Noted"
                                            })
                                update_row("tm_minutes", m["id"], {
                                    "title": etitle, "meeting_date": str(edate),
                                    "chair": echair, "location": eloc,
                                    "attendees": eatt, "apologies": eapol,
                                    "digital_summary": edigital,
                                    "digital_items": json.dumps(parsed_items),
                                    "face_to_face_content": eff,
                                    "content": econtent,
                                    "action_summary": eactions
                                })
                                del st.session_state[f"editing_min_{m['id']}"]
                                st.rerun()
    else:
        st.info("No minutes recorded yet. Use the form below to add the first entry.")

    # ── NEW MINUTES FORM (admin only) ─────────────────────────────────────────
    if not st.session_state.get("admin"):
        return

    st.divider()

    if is_staff:
        _render_new_staff_minutes()
    else:
        _render_new_team_minutes(program)


# ─────────────────────────────────────────────────────────────────────────────

def _render_new_staff_minutes():
    st.subheader("✍️ Record Combined Staff Meeting Minutes")
    st.markdown("""
    <div style="background:#f0f9ff;border:1px solid #bae6fd;border-radius:8px;
                padding:10px 16px;margin-bottom:16px;font-size:13px;color:#0c4a6e;">
      Enter meeting details, attendance, then complete <strong>Part 1 (digital)</strong>
      and <strong>Part 2 (face-to-face)</strong>. Everything saves as one unified record
      with a downloadable Word document.
    </div>""", unsafe_allow_html=True)

    with st.form("new_staff_minutes_form", clear_on_submit=True):

        # Details
        st.markdown("##### 📋 Meeting Details")
        dc1, dc2, dc3 = st.columns([3, 2, 2])
        with dc1: m_title    = st.text_input("Meeting title *", placeholder="e.g. Staff Meeting — T1 W4")
        with dc2: m_date     = st.date_input("Date *", value=datetime.date.today())
        with dc3: m_chair    = st.text_input("Chair", placeholder="e.g. Candice Cooper")
        m_location = st.text_input("Location / format", placeholder="e.g. Staff Room / Teams")

        # Attendance
        st.markdown("##### 🙋 Attendance")
        ac1, ac2 = st.columns(2)
        with ac1:
            m_attendees = st.text_area("Present — one name per line",
                placeholder="Candice Cooper\nJane Smith\nBob Jones", height=100)
        with ac2:
            m_apologies = st.text_area("Apologies — one name per line",
                placeholder="Sam Brown", height=100)

        # Part 1 Digital
        st.markdown("##### 💻 Part 1 — Digital Staff Meeting")
        m_digital_summary = st.text_area(
            "Digital meeting summary",
            placeholder="Brief overview of what staff were asked to review / action online this week. "
                        "E.g. Staff reviewed 3 notices: timetable update, PD day reminder, transport change.",
            height=90
        )
        st.markdown(
            "**Items table** — one item per line in format: "
            "`Item description | Raised By | Actioned By | Status`"
        )
        st.caption("Status options: Actioned · Noted · Pending")
        m_digital_items_raw = st.text_area(
            "Digital items",
            placeholder=(
                "Timetable update Week 5 | Candice Cooper | All staff | Noted\n"
                "Student transport change JP | Admin | JP team | Actioned\n"
                "PD Day reminder 14 March | Candice Cooper | All staff | Actioned"
            ),
            height=110
        )

        # Part 2 Face-to-face
        st.markdown("##### 👥 Part 2 — Face-to-Face Meeting")
        ff_method = st.radio(
            "Minutes input method",
            ["📄 Type / paste minutes directly",
             "🎙️ Drop transcript — AI will structure and improve it"],
            horizontal=True
        )
        do_ai = "🎙️" in ff_method
        m_ff_content = st.text_area(
            "Face-to-face minutes" if not do_ai else "Paste transcript or rough notes",
            height=280,
            placeholder=(
                "Type your face-to-face meeting minutes here..."
                if not do_ai else
                "Paste the recorded transcript or rough notes. AI will turn this into polished minutes."
            )
        )

        # Actions
        st.markdown("##### ✅ Action Items")
        st.caption("One per line: `Action | Assigned To | Due date`")
        m_actions = st.text_area(
            "Action items",
            placeholder="Update student profiles | Candice Cooper | Week 5\nBook PD venue | Admin | 14 March",
            height=80
        )

        submitted = st.form_submit_button(
            "✨ Improve with AI & Save" if do_ai else "💾 Save Combined Minutes",
            type="primary"
        )

        if submitted:
            if not m_title.strip():
                st.warning("Please enter a meeting title.")
            else:
                attendees_str = ", ".join(
                    n.strip() for n in m_attendees.replace("\n",",").split(",") if n.strip()
                )
                apologies_str = ", ".join(
                    n.strip() for n in m_apologies.replace("\n",",").split(",") if n.strip()
                )
                digital_items = []
                for line in m_digital_items_raw.splitlines():
                    parts = [p.strip() for p in line.split("|")]
                    if parts and parts[0]:
                        digital_items.append({
                            "item":        parts[0],
                            "raised_by":   parts[1] if len(parts) > 1 else "",
                            "actioned_by": parts[2] if len(parts) > 2 else "",
                            "status":      parts[3] if len(parts) > 3 else "Noted",
                        })
                final_ff = m_ff_content.strip()
                if do_ai and final_ff:
                    with st.spinner("AI is improving your minutes…"):
                        final_ff = improve_with_ai(final_ff, "STAFF", str(m_date))

                insert("tm_minutes", {
                    "program":              "STAFF",
                    "title":               m_title.strip(),
                    "meeting_date":        str(m_date),
                    "chair":               m_chair.strip(),
                    "location":            m_location.strip(),
                    "attendees":           attendees_str,
                    "apologies":           apologies_str,
                    "digital_summary":     m_digital_summary.strip(),
                    "digital_items":       json.dumps(digital_items),
                    "face_to_face_content": final_ff,
                    "content":             "",
                    "action_summary":      m_actions.strip(),
                })
                st.success("✅ Combined meeting minutes saved!")
                st.rerun()


def _render_new_team_minutes(program):
    st.subheader("✍️ Record New Minutes")

    ff_method = st.radio(
        "Minutes input method",
        ["📄 Type / paste minutes directly",
         "🎙️ Drop transcript — AI will structure and improve it"],
        horizontal=True,
        key=f"ff_method_{program}"
    )
    do_ai = "🎙️" in ff_method

    with st.form(f"new_minutes_form_{program}", clear_on_submit=True):

        st.markdown("##### 📋 Meeting Details")
        dc1, dc2, dc3 = st.columns([3, 2, 2])
        with dc1: m_title    = st.text_input("Meeting title *", placeholder=f"{PROGRAMS[program]['label']} Meeting — T1 W4")
        with dc2: m_date     = st.date_input("Date *", value=datetime.date.today())
        with dc3: m_chair    = st.text_input("Chair")
        m_location = st.text_input("Location / format", placeholder="e.g. Room 3 / Teams")

        st.markdown("##### 🙋 Attendance")
        ac1, ac2 = st.columns(2)
        with ac1: m_attendees = st.text_area("Present — one name per line", height=90)
        with ac2: m_apologies = st.text_area("Apologies — one name per line", height=90)

        schedules  = fetch("tm_schedules", {"program": program})
        sched_opts = {s["schedule_name"]: s["id"] for s in schedules}
        sched_opts["General / Unassigned"] = None
        m_sched = st.selectbox("Associated schedule", list(sched_opts.keys()))

        st.markdown("##### 📝 Minutes")
        m_content = st.text_area(
            "Minutes" if not do_ai else "Paste transcript or rough notes",
            height=280,
            placeholder=(
                "Type your meeting minutes here..."
                if not do_ai else
                "Paste the recorded transcript or rough notes. AI will turn this into polished minutes."
            )
        )

        st.markdown("##### ✅ Action Items")
        st.caption("One per line: `Action | Assigned To | Due date`")
        m_actions = st.text_area(
            "Action items",
            placeholder="Update student profiles | Jane Smith | Week 5",
            height=80
        )

        submitted = st.form_submit_button(
            "✨ Improve with AI & Save" if do_ai else "💾 Save Minutes",
            type="primary"
        )

        if submitted:
            if not m_title.strip():
                st.warning("Please enter a meeting title.")
            else:
                attendees_str = ", ".join(
                    n.strip() for n in m_attendees.replace("\n",",").split(",") if n.strip()
                )
                apologies_str = ", ".join(
                    n.strip() for n in m_apologies.replace("\n",",").split(",") if n.strip()
                )
                final_content = m_content.strip()
                if do_ai and final_content:
                    with st.spinner("AI is improving your minutes…"):
                        final_content = improve_with_ai(final_content, program, str(m_date))
                insert("tm_minutes", {
                    "program":              program,
                    "schedule_id":          sched_opts.get(m_sched),
                    "title":               m_title.strip(),
                    "meeting_date":        str(m_date),
                    "chair":               m_chair.strip(),
                    "location":            m_location.strip(),
                    "attendees":           attendees_str,
                    "apologies":           apologies_str,
                    "content":             final_content,
                    "action_summary":      m_actions.strip(),
                    "digital_summary":     "",
                    "digital_items":       "[]",
                    "face_to_face_content": "",
                })
                st.success("✅ Minutes saved!")
                st.rerun()

def render_actions(program):
    """Action items tracker."""
    st.subheader("✅ Action Items")

    actions = fetch("tm_actions", {"program": program}, "due_date")

    open_actions = [a for a in actions if a.get("status") != "completed"]
    completed    = [a for a in actions if a.get("status") == "completed"]

    if open_actions:
        st.markdown("**Open Actions**")
        for a in open_actions:
            c1, c2, c3 = st.columns([4, 2, 1])
            with c1:
                st.markdown(f"🔲 **{a['action']}**")
                st.caption(f"Assigned to: {a.get('assigned_to','—')} · Due: {a.get('due_date','—')}")
            with c2:
                if st.session_state.get("admin"):
                    if st.button("✅ Mark complete", key=f"complete_{a['id']}"):
                        update_row("tm_actions", a["id"], {"status": "completed"})
                        st.rerun()
            with c3:
                if st.session_state.get("admin"):
                    if st.button("🗑️", key=f"del_action_{a['id']}"):
                        delete_row("tm_actions", a["id"])
                        st.rerun()
            st.divider()
    else:
        st.info("No open action items.")

    if completed:
        with st.expander(f"✅ Completed ({len(completed)})"):
            for a in completed:
                st.markdown(f"~~{a['action']}~~ — {a.get('assigned_to','—')}")

    if st.session_state.get("admin"):
        with st.expander("➕ Add Action Item"):
            with st.form("action_form"):
                a_text = st.text_input("Action")
                a_col1, a_col2 = st.columns(2)
                with a_col1:
                    a_assigned = st.text_input("Assigned to")
                with a_col2:
                    a_due = st.date_input("Due date", value=datetime.date.today() + datetime.timedelta(weeks=2))
                a_notes = st.text_input("Notes (optional)")
                if st.form_submit_button("Add"):
                    if a_text.strip():
                        insert("tm_actions", {
                            "program": program,
                            "action": a_text.strip(),
                            "assigned_to": a_assigned.strip(),
                            "due_date": str(a_due),
                            "notes": a_notes.strip(),
                            "status": "open"
                        })
                        st.success("Action added!")
                        st.rerun()


def render_attendance(program):
    """Attendance register."""
    st.subheader("🙋 Attendance")

    schedules = fetch("tm_schedules", {"program": program})
    if not schedules:
        st.info("Add a meeting schedule first.")
        return

    sched_opts = {s["schedule_name"]: s["id"] for s in schedules}
    sel_sched = st.selectbox("Meeting", list(sched_opts.keys()), key="att_sched")
    sel_date  = st.date_input("Meeting date", value=datetime.date.today(), key="att_date")

    existing = supabase.table("tm_attendance").select("*").eq("program", program)\
        .eq("schedule_id", sched_opts[sel_sched]).eq("meeting_date", str(sel_date)).execute().data or []

    if existing:
        st.markdown("**Attendance record:**")
        for att in existing:
            status_icon = "✅" if att["status"] == "present" else "❌"
            st.markdown(f"{status_icon} {att['staff_name']} — {att['status'].title()}")
    else:
        st.info("No attendance recorded for this meeting date yet.")

    if st.session_state.get("admin"):
        with st.expander("📝 Record Attendance"):
            with st.form("attendance_form"):
                names_raw = st.text_area("Staff names (one per line)")
                status_val = st.selectbox("Status for all entered", ["present", "absent", "apology"])
                if st.form_submit_button("Save"):
                    names = [n.strip() for n in names_raw.splitlines() if n.strip()]
                    for name in names:
                        # Upsert by checking if already exists
                        check = supabase.table("tm_attendance").select("id")\
                            .eq("program", program).eq("schedule_id", sched_opts[sel_sched])\
                            .eq("meeting_date", str(sel_date)).eq("staff_name", name).execute().data
                        if check:
                            update_row("tm_attendance", check[0]["id"], {"status": status_val})
                        else:
                            insert("tm_attendance", {
                                "program": program,
                                "schedule_id": sched_opts[sel_sched],
                                "meeting_date": str(sel_date),
                                "staff_name": name,
                                "status": status_val
                            })
                    st.success(f"Attendance saved for {len(names)} staff member(s).")
                    st.rerun()


def render_documents(program):
    """Supporting documents — links/notes."""
    st.subheader("📎 Supporting Documents")

    docs = fetch("tm_documents", {"program": program}, "created_at")

    if docs:
        for d in docs:
            c1, c2 = st.columns([5, 1])
            with c1:
                if d.get("url"):
                    st.markdown(f"📄 [{d['title']}]({d['url']})")
                else:
                    st.markdown(f"📄 **{d['title']}**")
                if d.get("description"):
                    st.caption(d["description"])
                st.caption(f"Added: {d.get('created_at','')[:10]}")
            if st.session_state.get("admin"):
                with c2:
                    if st.button("🗑️", key=f"del_doc_{d['id']}"):
                        delete_row("tm_documents", d["id"])
                        st.rerun()
            st.divider()
    else:
        st.info("No documents added yet.")

    if st.session_state.get("admin"):
        with st.expander("➕ Add Document Link"):
            with st.form("doc_form"):
                d_title = st.text_input("Title")
                d_url   = st.text_input("URL (optional)")
                d_desc  = st.text_area("Description (optional)")
                if st.form_submit_button("Add"):
                    if d_title.strip():
                        insert("tm_documents", {
                            "program": program, "title": d_title.strip(),
                            "url": d_url.strip(), "description": d_desc.strip()
                        })
                        st.success("Document added!")
                        st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
def main():
    program = get_program_from_params()
    prog_info = PROGRAMS[program]
    week_range = get_week_from_params()

    admin_login()

    # ── Header ──
    grad_end = {
        "JP": "#3fa86a", "PY": "#2a6fd4", "SY": "#9b59b6",
        "STAFF": "#2a9fa8", "ADMIN": "#c97b1a"
    }.get(program, "#2a6fd4")

    st.markdown(f"""
    <div style="background:linear-gradient(135deg,{prog_info['color']},{grad_end});
                padding:1.5rem 2rem;border-radius:16px;margin-bottom:1.5rem;">
      <div style="font-size:2.5rem;margin-bottom:0.25rem;">{prog_info['emoji']}</div>
      <h1 style="color:white;margin:0;font-size:1.8rem;">{prog_info['label']}</h1>
      <p style="color:rgba(255,255,255,0.85);margin:0.25rem 0 0;">Cowandilla Learning Centre · LBU</p>
    </div>
    """, unsafe_allow_html=True)

    # ── Program switcher ──
    with st.sidebar:
        st.markdown("### 🔀 Switch Program")
        for p, info in PROGRAMS.items():
            if p != program:
                st.markdown(f"[{info['emoji']} {info['label']}](?program={p})")
        st.divider()
        if st.session_state.get("admin"):
            st.markdown("### ⚙️ Admin Mode Active")
            st.caption("You can add, edit and delete all content.")

    # ── Tabs ──
    tab_sched, tab_agenda, tab_minutes, tab_actions, tab_attend, tab_docs = st.tabs([
        "📅 Schedules", "📋 Agenda", "📝 Minutes", "✅ Actions", "🙋 Attendance", "📎 Documents"
    ])

    with tab_sched:
        render_schedules(program)
    with tab_agenda:
        render_agenda(program)
    with tab_minutes:
        render_minutes(program, week_range=week_range if program == "STAFF" else None)
    with tab_actions:
        render_actions(program)
    with tab_attend:
        render_attendance(program)
    with tab_docs:
        render_documents(program)


if __name__ == "__main__":
    main()
